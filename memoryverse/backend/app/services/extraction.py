"""
MemoryVerse AI — Text Extraction Service

Extracts raw text from uploaded documents (PDF, DOCX, images).
Uses PyMuPDF for PDFs, python-docx for DOCX, pytesseract for images.
Per Rules.md §3: every extraction call is wrapped in specific error handling.

Original files are never modified — all operations are read-only.
"""

import logging
from pathlib import Path

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Supported MIME types mapped to extraction method
SUPPORTED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "image/png": "image",
    "image/jpeg": "image",
    "image/jpg": "image",
}


def extract_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using PyMuPDF.

    Args:
        file_path: Absolute path to the PDF file.

    Returns:
        Extracted text content as a string.

    Raises:
        ExtractionError: If the PDF cannot be read or contains no extractable text.
    """
    try:
        doc = fitz.open(file_path)
        text_parts = []
        for page_num, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)
        doc.close()

        full_text = "\n\n".join(text_parts).strip()

        if not full_text:
            raise ExtractionError(
                message="PDF contains no extractable text",
                detail=f"File '{Path(file_path).name}' may be a scanned document without embedded text.",
                suggestion="Try uploading the document as an image (PNG/JPG) so OCR can process it, "
                           "or re-save the PDF with text layer enabled.",
            )

        logger.info(
            "Extracted %d characters from PDF '%s' (%d pages)",
            len(full_text), Path(file_path).name, len(text_parts),
        )
        return full_text

    except ExtractionError:
        raise
    except Exception as e:
        logger.error("PDF extraction failed for '%s': %s", file_path, str(e))
        raise ExtractionError(
            message="Failed to read PDF file",
            detail=str(e),
            suggestion="The file may be corrupted or password-protected. "
                       "Try re-saving it without protection and uploading again.",
        )


def extract_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        file_path: Absolute path to the DOCX file.

    Returns:
        Extracted text content as a string.

    Raises:
        ExtractionError: If the DOCX cannot be read or is empty.
    """
    try:
        doc = DocxDocument(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts).strip()

        if not full_text:
            raise ExtractionError(
                message="DOCX file contains no extractable text",
                detail=f"File '{Path(file_path).name}' appears to be empty.",
                suggestion="Check that the document contains text content and try again.",
            )

        logger.info(
            "Extracted %d characters from DOCX '%s'",
            len(full_text), Path(file_path).name,
        )
        return full_text

    except ExtractionError:
        raise
    except Exception as e:
        logger.error("DOCX extraction failed for '%s': %s", file_path, str(e))
        raise ExtractionError(
            message="Failed to read DOCX file",
            detail=str(e),
            suggestion="The file may be corrupted. Try re-saving it in Word and uploading again.",
        )


def extract_from_image(file_path: str, tesseract_cmd: str = "") -> str:
    """
    Extract text from an image file using Tesseract OCR.

    Args:
        file_path: Absolute path to the image file.
        tesseract_cmd: Optional path to Tesseract executable.

    Returns:
        Extracted text content as a string.

    Raises:
        ExtractionError: If OCR fails or produces no text.
    """
    try:
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

        image = Image.open(file_path)
        text = pytesseract.image_to_string(image).strip()

        if not text:
            raise ExtractionError(
                message="OCR could not extract text from image",
                detail=f"File '{Path(file_path).name}' produced no readable text.",
                suggestion="Ensure the image is clear, well-lit, and contains readable text. "
                           "Try uploading a higher-resolution version.",
            )

        logger.info(
            "Extracted %d characters via OCR from image '%s'",
            len(text), Path(file_path).name,
        )
        return text

    except ExtractionError:
        raise
    except Exception as e:
        logger.error("OCR extraction failed for '%s': %s", file_path, str(e))
        raise ExtractionError(
            message="OCR processing failed",
            detail=str(e),
            suggestion="Make sure Tesseract OCR is installed on your system. "
                       "On Windows, install from https://github.com/UB-Mannheim/tesseract/wiki",
        )


def extract_text(file_path: str, mime_type: str, tesseract_cmd: str = "") -> str:
    """
    Dispatch text extraction to the appropriate handler based on MIME type.

    Args:
        file_path: Absolute path to the uploaded file.
        mime_type: MIME type of the file.
        tesseract_cmd: Optional path to Tesseract executable.

    Returns:
        Extracted text content as a string.

    Raises:
        ExtractionError: If the file type is unsupported or extraction fails.
    """
    extraction_type = SUPPORTED_TYPES.get(mime_type)

    if extraction_type is None:
        raise ExtractionError(
            message=f"Unsupported file type: {mime_type}",
            detail=f"Supported types: PDF, DOCX, PNG, JPG/JPEG.",
            suggestion="Convert your document to one of the supported formats and try again.",
        )

    if extraction_type == "pdf":
        return extract_from_pdf(file_path)
    elif extraction_type == "docx":
        return extract_from_docx(file_path)
    elif extraction_type == "image":
        return extract_from_image(file_path, tesseract_cmd)
    else:
        raise ExtractionError(
            message=f"No extraction handler for type: {extraction_type}",
            detail="This is a bug — please report it.",
            suggestion="Try uploading the file in a different format.",
        )


class ExtractionError(Exception):
    """
    Custom exception for text extraction failures.
    Carries user-facing message, technical detail, and actionable suggestion
    per Rules.md §3.
    """

    def __init__(self, message: str, detail: str = "", suggestion: str = ""):
        self.message = message
        self.detail = detail
        self.suggestion = suggestion
        super().__init__(message)
